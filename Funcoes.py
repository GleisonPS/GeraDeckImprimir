import os
from tkinter import *
from tkinter import messagebox,filedialog
import urllib.request as request
from docx import Document
from docx.shared import Inches,Mm
from docx.enum.section import WD_ORIENT

class PlaceholderEntry(Entry):
    def __init__(self, master=None, placeholder="", *args, **kwargs):
        super().__init__(master, *args, **kwargs)
        self.placeholder = placeholder
        self.bind("<FocusIn>", self._on_focus_in)
        self.bind("<FocusOut>", self._on_focus_out)
        self._set_placeholder()

    def _set_placeholder(self):
        self.insert(0, self.placeholder)
        self.config(fg='black', font=('Arial', 15))

    def _on_focus_in(self, event):
        if self.get() == self.placeholder:
            self.delete(0, END)
            self.config(fg='black', font=('Arial', 15))

    def _on_focus_out(self, event):
        if self.get() == '':
            self._set_placeholder()

def abrir_pasta(choose_file_path,nome_arquivo):
    caminho = upload_file()
    if caminho:
        choose_file_path.delete(0, 'end')
        choose_file_path.insert(0, caminho)

        # Extrai o nome do arquivo sem a extensão e atualiza a entrada `nome_arquivo`
        file_name = os.path.splitext(os.path.basename(caminho))[0]
        nome_arquivo.delete(0, 'end')  # Limpa a entrada
        nome_arquivo.insert(0, file_name)
    
def criar_deck(choose_file_path,nome_arquivo):
    rota = choose_file_path.get()
    nome_do_arquivo = nome_arquivo.get()
    if rota and nome_do_arquivo:
        Init_Baixar(rota, nome_do_arquivo)
    else:
        messagebox.showwarning('Dados incompletos', 'Por favor, selecione um arquivo e um nome válidos.')


# Função que é chamada quando o botão de upload é clicado
def upload_file():
    # Abre uma caixa de diálogo para seleção de arquivo
    file_path = filedialog.askopenfilename(
        title="Selecione um arquivo",
        filetypes=[
            ("Arquivos YDK", "*.ydk"),
            ("Arquivos de Texto", "*.txt")
        ]
    )
    
    if file_path:  # Verifica se um arquivo foi selecionado
        #messagebox.showinfo("Arquivo Selecionado", f"O arquivo {file_path} foi selecionado!")
        return file_path
    else:
        messagebox.showwarning("Nenhum arquivo", "Nenhum arquivo foi selecionado.")


def baixar(card):
    url = "https://images.ygoprodeck.com/images/cards/"
    pasta = "Cards"
    
    # Verifica se a pasta existe, caso contrário, cria
    if not os.path.exists(pasta):
        os.makedirs(pasta)
    
    try:
        if not os.path.isfile(f"Cards/{card}.jpg"):
            request.urlretrieve(f"{url}{card}.jpg", f"{pasta}/{card}.jpg")
    except Exception as ex:
        print(f"Erro ao baixar a imagem: {ex}")


def Init_Baixar(Arquivo,nome="teste"):
    cards = []
    documento = Document()

    #Configurando as margens do documento como 0
    section = documento.sections[0]
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    documento.element

    # Modifica as propriedades de página para paisagem
    section.orientation = WD_ORIENT.LANDSCAPE
    # Ajusta a largura e altura da página para a orientação paisagem
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height



    with open(Arquivo, "r") as arquivo:
        for linha in arquivo:
            if linha.strip() in ["#created by Toupeira","#main","#extra","!side"]:
                pass
            else:
                cards.append(linha.strip())

    paragraph = documento.add_paragraph()

    # Adiciona a imagem no mesmo parágrafo
    run = paragraph.add_run()
    
    for codeC in cards:
        baixar(codeC)
        try:
            Add_img(run,codeC)
        except:
            pass

    directory = "deckPronto"
    if not os.path.exists(directory):
        os.makedirs(directory)


    documento.save(f"deckPronto/{nome}.docx")
    messagebox.showinfo("Arquivo Criado", f"O arquivo foi Criado!")


def Add_img(documento,codeCard):
    
    caminho = f"Cards/{codeCard}.jpg"
    
    # Converte as dimensões de mm para polegadas
    largura_polegadas = 57 / 25.4
    altura_polegadas = 89 / 25.4

    documento.add_picture(caminho, width=Inches(largura_polegadas), height=Inches(altura_polegadas))
    
    #Adicionar um espaço depois da imagem
    documento.add_text("    ")
